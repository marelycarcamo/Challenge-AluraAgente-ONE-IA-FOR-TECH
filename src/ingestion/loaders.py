# 📖 ¿Cómo funciona load_pdf()?

# Objetivo: leer un documento en formato PDF y convertirlo en un objeto Document, que será el formato estándar utilizado por el pipeline RAG.

# Proceso de extracción
# Verifica que el archivo exista en la ruta indicada.
# Abre el documento utilizando la librería PyPDF (PdfReader).
# Recorre el PDF página por página.
# En cada iteración intenta extraer el texto mediante extract_text().
# El texto obtenido se va acumulando en una única variable (text) que contendrá el contenido completo del documento.
# Finalmente, construye y retorna un objeto Document con:
# Título: nombre del archivo sin extensión.
# Origen: ruta del archivo.
# Extensión: tipo de archivo (pdf).
# Contenido: texto completo extraído del documento.
# Metadatos: diccionario preparado para almacenar información adicional (por ejemplo, número de páginas, fecha de extracción o departamento municipal).

"""
Module: loaders.py

Purpose
-------
Load documents from different file formats (PDF, DOCX and HTML)
and convert them into the common Document model used by the RAG pipeline.

Responsibilities
----------------
- Read files from disk.
- Extract textual content.
- Create and return Document objects.

This module DOES NOT:
- Clean text.
- Normalize text.
- Split text into chunks.
- Generate embeddings.
"""

from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from src.models.document import Document
from uuid import uuid4


def load_pdf(file_path: Path) -> Document:
    """
    Load a PDF document and convert it into a Document object.

    Args:
        file_path (Path): Path to the PDF file.

    Returns:
        Document: Document containing the extracted text.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the PDF cannot be processed.
    """

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        reader = PdfReader(file_path)

        text = ""

        for page in reader.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

        return Document(
            id=str(uuid4()),
            title=file_path.stem,
            source=str(file_path),
            file_type=file_path.suffix.lower().replace(".", ""),
            content=text,
            metadata={
                "pages": len(reader.pages),
            }
        )
            

    except Exception as e:
        raise ValueError(f"Error loading PDF '{file_path}': {e}")


def load_docx(file_path: Path) -> Document:
    """
    Load a Word (.docx) document and convert it into a Document object.
    """

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        doc = DocxDocument(file_path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        )

        return Document(
            id=str(uuid4()),
            title=file_path.stem,
            source=str(file_path),
            file_type=file_path.suffix.lower().replace(".", ""),
            content=text,
            metadata={
                "paragraphs": len(doc.paragraphs),
            }
        )

    except Exception as e:
        raise ValueError(f"Error loading DOCX '{file_path}': {e}")


def load_html(file_path: Path) -> Document:
    """
    Load an HTML document and convert it into a Document object.
    """

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        html = file_path.read_text(encoding="utf-8")

        soup = BeautifulSoup(html, "html.parser")

        text = soup.get_text(separator="\n", strip=True)

        return Document(
            id=str(uuid4()),
            title=file_path.stem,
            source=str(file_path),
            file_type=file_path.suffix.lower().replace(".", ""),
            content=text,
            metadata={
                "title": soup.title.string if soup.title else None,
                "h1": len(soup.find_all("h1")),
                "h2": len(soup.find_all("h2"))               
            }
        )

    except Exception as e:
        raise ValueError(f"Error loading HTML '{file_path}': {e}")


def load_file(file_path: Path) -> Document:
    """
    Detect the file type and call the appropriate loader.

    Args:
        file_path (Path): Path to the document.

    Returns:
        Document

    Raises:
        ValueError: If the file type is not supported.
    """

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return load_pdf(file_path)

    if extension == ".docx":
        return load_docx(file_path)

    if extension in [".html", ".htm"]:
        return load_html(file_path)

    raise ValueError(f"Unsupported file type: {extension}")
